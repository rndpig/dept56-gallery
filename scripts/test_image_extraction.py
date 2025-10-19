"""Test script to explore image extraction from Word documents"""
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import sys

if len(sys.argv) < 2:
    print("Usage: python test_image_extraction.py <path_to_docx>")
    sys.exit(1)

docx_path = sys.argv[1]
doc = Document(docx_path)

print(f"Document: {docx_path}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total sections: {len(doc.sections)}")
print()

# Method 1: Check inline shapes in paragraphs
print("=" * 70)
print("METHOD 1: Inline Shapes in Paragraphs")
print("=" * 70)
for i, para in enumerate(doc.paragraphs):
    if para._element.xpath('.//pic:pic'):
        print(f"Paragraph {i}: Contains inline image")
        # Try to get image details
        for run in para.runs:
            if 'graphicData' in run._element.xml:
                print(f"  - Run contains graphic data")

# Method 2: Check document relationships for images
print()
print("=" * 70)
print("METHOD 2: Document Part Relationships")
print("=" * 70)
image_parts = []
for rel in doc.part.rels.values():
    if "image" in rel.target_ref:
        print(f"Found image: {rel.target_ref}")
        image_parts.append(rel)

print(f"\nTotal images in document: {len(image_parts)}")

# Method 3: Try to associate images with page/section
print()
print("=" * 70)
print("METHOD 3: Image Locations by Paragraph")
print("=" * 70)

# Track current page and find images
current_page = 1
for i, para in enumerate(doc.paragraphs):
    # Check for page break
    if para._element.xpath('.//w:br[@w:type="page"]'):
        current_page += 1
        print(f"\n>>> PAGE BREAK at paragraph {i} -> Now on page {current_page}\n")
    
    # Check for image
    if para._element.xpath('.//pic:pic'):
        print(f"📷 Image found at paragraph {i} (Page {current_page})")
        # Try to extract image
        for rel_id in para._element.xpath('.//a:blip/@r:embed'):
            rel = doc.part.rels[rel_id]
            print(f"   Image file: {rel.target_ref}")
            print(f"   Image bytes: {len(rel.target_part.blob)} bytes")

print(f"\n{'=' * 70}")
print(f"Summary: Found images on page(s) 1 and/or 2+")
print(f"{'=' * 70}")
