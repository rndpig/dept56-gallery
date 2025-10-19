"""
Debug script to inspect Word document structure
"""

from docx import Document
import sys

if len(sys.argv) < 2:
    print("Usage: python debug_docx.py <path_to_docx>")
    sys.exit(1)

docx_path = sys.argv[1]
doc = Document(docx_path)

print(f"\n{'='*70}")
print(f"DOCUMENT STRUCTURE DEBUG")
print(f"{'='*70}\n")

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total sections: {len(doc.sections)}")
print(f"Total tables: {len(doc.tables)}")

print(f"\n{'─'*70}")
print("FIRST 20 PARAGRAPHS:")
print(f"{'─'*70}\n")

for i, para in enumerate(doc.paragraphs[:20]):
    text = para.text.strip()
    if text:
        print(f"{i:2d}. {text[:100]}")
    else:
        print(f"{i:2d}. (empty)")

print(f"\n{'─'*70}")
print("PAGE BREAKS:")
print(f"{'─'*70}\n")

page_break_count = 0
for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
            page_break_count += 1
            print(f"  Page break found at paragraph {i}")

print(f"\nTotal page breaks: {page_break_count}")

print(f"\n{'─'*70}")
print("TABLES:")
print(f"{'─'*70}\n")

for i, table in enumerate(doc.tables):
    print(f"\nTable {i+1}:")
    for row_idx, row in enumerate(table.rows[:3]):  # First 3 rows
        cells_text = [cell.text.strip()[:30] for cell in row.cells]
        print(f"  Row {row_idx}: {' | '.join(cells_text)}")

print()
