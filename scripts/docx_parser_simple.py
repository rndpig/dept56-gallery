"""
Simplified Word Document Parser for Department 56 Gallery
Extracts houses and accessories from Word documents

Author: GitHub Copilot
Date: October 18, 2025
"""

import os
import re
import json
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict
from docx import Document

@dataclass
class ExtractedItem:
    """Represents an extracted house or accessory"""
    name: str
    item_type: str  # 'house' or 'accessory'
    sku: Optional[str] = None
    description: Optional[str] = None
    notes: Optional[str] = None
    price: Optional[float] = None
    year: Optional[int] = None  # Released/Introduced year
    retired_year: Optional[int] = None  # Retired year
    purchased_year: Optional[int] = None
    collection: Optional[str] = None  # Collection name
    series: Optional[str] = None  # Series name
    linked_items: List[str] = None
    image_path: Optional[str] = None  # Path to extracted image file
    
    def __post_init__(self):
        if self.linked_items is None:
            self.linked_items = []


class Dept56DocxParser:
    """
    Parser for Department 56 Word documents
    Expected format:
    - Line 1: House name
    - Line 2: Accessory name
    - Line 3: Box info (optional, ignored)
    - Following lines: Details (SKU, price, description, etc.)
    - Then accessory details
    """
    
    def __init__(self, output_dir: str = "parsed_output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def parse_document(self, docx_path: str) -> Dict:
        """Parse a Word document and extract houses and accessories"""
        print(f"\n{'='*70}")
        print(f"Parsing: {Path(docx_path).name}")
        print(f"{'='*70}\n")
        
        doc = Document(docx_path)
        
        # Extract all non-empty paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        print(f"Found {len(paragraphs)} text paragraphs\n")
        
        # Parse the content
        house, accessories = self._parse_content(paragraphs)
        
        # Extract images from the document
        house, accessories = self._extract_images(doc, house, accessories, docx_path)
        
        result = {
            'source_file': Path(docx_path).name,
            'houses': [asdict(house)] if house else [],
            'accessories': [asdict(a) for a in accessories],
        }
        
        # Print summary
        print(f"\n{'─'*70}")
        print(f"Summary:")
        print(f"  Houses: {1 if house else 0}")
        print(f"  Accessories: {len(accessories)}")
        print(f"{'─'*70}\n")
        
        return result
    
    def _parse_content(self, paragraphs: List[str]):
        """Parse the paragraphs and extract house and accessories"""
        if len(paragraphs) < 2:
            print("⚠ Warning: Not enough paragraphs")
            return None, []
        
        # Line 0: House name
        house_name = paragraphs[0]
        print(f"🏠 House name: {house_name}")
        
        # Lines 1+: Might be one or more accessory names, OR "No ACC Box", OR "Box X"
        # Collect potential accessory names until we hit a box line or house details
        canonical_accessory_names = []
        start_idx = 1
        
        for i in range(1, min(len(paragraphs), 5)):  # Check up to 4 lines after house name
            line = paragraphs[i]
            
            # Stop if we hit a box line
            if self._is_box_line(line):
                print(f"   Skipping box line: {line}")
                start_idx = i + 1
                break
            
            # Stop if it's a "No ACC Box" marker
            if self._is_no_accessory_marker(line):
                print(f"   ℹ️  No traditional accessory found (marker: '{line}')")
                start_idx = i + 1
                break
            
            # Stop if it's just "Box X"
            if re.match(r'^\s*Box\s+\d+\s*$', line, re.IGNORECASE):
                print(f"   ℹ️  Line {i+1} is just a box reference, no more accessory names")
                start_idx = i + 1
                break
            
            # Stop if line looks like house metadata (SKU, price, description start)
            if any(keyword in line.lower() for keyword in ['sku', 'item #', '$', 'introduced', 'retired']):
                start_idx = i
                break
            
            # Otherwise, treat as accessory name
            canonical_accessory_names.append(line)
            print(f"📦 Accessory {len(canonical_accessory_names)} name: {line}")
            start_idx = i + 1
        
        # Look for "Accessories-..." line anywhere in the document
        accessory_list_line = None
        accessory_list_idx = None
        for i in range(start_idx, len(paragraphs)):
            if re.match(r'^\s*Accessor(?:y|ies)\s*[-:]', paragraphs[i], re.IGNORECASE):
                accessory_list_line = paragraphs[i]
                accessory_list_idx = i
                print(f"   Found accessories list at line {i}: {accessory_list_line[:80]}")
                break
        
        # Find all accessory sections based on canonical names
        accessory_sections = {}  # Maps canonical name -> start index
        
        # Determine house content lines
        if accessory_list_idx is not None:
            # Include all lines EXCEPT the accessory list line
            # Combine lines before and after the accessory list
            house_lines = paragraphs[start_idx:accessory_list_idx] + paragraphs[accessory_list_idx + 1:]
        elif canonical_accessory_names:
            # Find where each accessory section begins
            for i in range(start_idx, len(paragraphs)):
                # Skip if line starts with "coordinates" or similar linking words
                if re.match(r'^\s*Coordinates\s+', paragraphs[i], re.IGNORECASE):
                    continue
                
                # Check if this line matches any canonical accessory name
                for canonical_name in canonical_accessory_names:
                    if self._names_match(paragraphs[i], canonical_name):
                        accessory_sections[canonical_name] = i
                        print(f"   Found accessory section '{canonical_name}' starting at line {i}")
                        break
            
            # Parse house - ends at first accessory section
            if accessory_sections:
                first_accessory_idx = min(accessory_sections.values())
                house_lines = paragraphs[start_idx:first_accessory_idx]
            else:
                # If we can't find the splits, take first half as house
                mid = len(paragraphs) // 2
                house_lines = paragraphs[start_idx:mid]
        else:
            # No traditional accessory and no list - all lines are house
            house_lines = paragraphs[start_idx:]
        
        house = self._parse_house(house_name, canonical_accessory_names, house_lines)
        
        # Parse accessories
        accessories = []
        
        # If we found an accessories list line, extract accessories from it
        if accessory_list_line:
            accessory_items = self._extract_accessory_list(accessory_list_line)
            print(f"   Extracted {len(accessory_items)} accessories from list:")
            for acc_name, acc_description in accessory_items:
                print(f"     - {acc_name}")
                if acc_description:
                    print(f"       Description: {acc_description[:60]}...")
                # Create accessory record with name, description, and link to house
                accessory = ExtractedItem(
                    name=acc_name,
                    item_type='accessory',
                    description=acc_description,
                    linked_items=[house_name],
                )
                accessories.append(accessory)
        elif accessory_sections:
            # Traditional format: detailed accessory sections
            # Sort sections by line number to process in order
            sorted_sections = sorted(accessory_sections.items(), key=lambda x: x[1])
            
            for i, (canonical_name, section_start_idx) in enumerate(sorted_sections):
                # Determine where this section ends
                if i < len(sorted_sections) - 1:
                    # Ends at the start of the next section
                    section_end_idx = sorted_sections[i + 1][1]
                    accessory_lines = paragraphs[section_start_idx:section_end_idx]
                else:
                    # Last section goes to end of document
                    accessory_lines = paragraphs[section_start_idx:]
                
                # Parse this accessory with its canonical name
                accessory = self._parse_accessory(canonical_name, accessory_lines, house_name)
                if accessory:
                    accessories.append(accessory)
        
        # Check for "Set of X" contents list (alternative accessory format)
        if not accessories:
            set_idx, set_accessories = self._extract_set_contents(paragraphs, start_idx)
            if set_accessories:
                print(f"   Extracted {len(set_accessories)} accessories from set contents:")
                for acc_name in set_accessories:
                    print(f"     - {acc_name}")
                    accessory = ExtractedItem(
                        name=acc_name,
                        item_type='accessory',
                        linked_items=[house_name],
                    )
                    accessories.append(accessory)
        
        # Update house's linked items
        if accessories and house:
            house.linked_items = [acc.name for acc in accessories]
        
        return house, accessories
    
    def _parse_house(self, name: str, accessory_names: List[str], detail_lines: List[str]) -> ExtractedItem:
        """Parse house details"""
        details_text = ' '.join(detail_lines)
        
        print(f"\n   House details ({len(detail_lines)} lines):")
        for line in detail_lines[:5]:  # Show first 5 lines
            print(f"     - {line[:80]}")
        
        house = ExtractedItem(
            name=name,
            item_type='house',
            sku=self._extract_sku(details_text),
            price=self._extract_price(details_text),
            description=self._extract_description(detail_lines),
            year=self._extract_year(details_text),
            retired_year=self._extract_retired_year(details_text),
            purchased_year=self._extract_purchased_year(details_text),
            collection=self._extract_collection(details_text),
            series=self._extract_series(details_text),
            notes=self._extract_notes(detail_lines),
            linked_items=accessory_names,  # Already a list
        )
        
        return house
    
    def _parse_accessory(self, canonical_name: str, lines: List[str], house_name: str) -> Optional[ExtractedItem]:
        """Parse accessory details using the canonical name from line 2 of page 1"""
        if not lines:
            return None
        
        # Use canonical name instead of the name from the accessory section
        # (which might be a "Coordinates with..." line)
        name = canonical_name
        
        # All lines are details (don't skip first line as it might be coordinates)
        detail_lines = lines
        details_text = ' '.join(detail_lines)
        
        print(f"\n   Accessory details ({len(detail_lines)} lines):")
        for line in detail_lines[:5]:
            print(f"     - {line[:80]}")
        
        accessory = ExtractedItem(
            name=name,
            item_type='accessory',
            sku=self._extract_sku(details_text),
            price=self._extract_price(details_text),
            description=self._extract_description(detail_lines),
            year=self._extract_year(details_text),
            retired_year=self._extract_retired_year(details_text),
            purchased_year=self._extract_purchased_year(details_text),
            collection=self._extract_collection(details_text),
            series=self._extract_series(details_text),
            notes=self._extract_notes(detail_lines),
            linked_items=[house_name],  # Always link back to the house (1:1 relationship)
        )
        
        return accessory
    
    def _extract_sku(self, text: str) -> Optional[str]:
        """Extract SKU/Item number"""
        patterns = [
            r'SKU[\s:]+([A-Z0-9.-]+)',
            r'Item\s*#[\s:]*([A-Z0-9.-]+)',  # Handles "Item #56957" and "Item # 56957"
            r'\b(\d{2,3}\.\d{4,5})\b',  # e.g., 56.12345
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)
        return None
    
    def _extract_price(self, text: str) -> Optional[float]:
        """Extract price"""
        patterns = [
            r'\$\s*(\d+(?:\.\d{2})?)',
            r'(?:Price|Cost|MSRP)[\s:]*\$?\s*(\d+(?:\.\d{2})?)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return float(match.group(1))
                except ValueError:
                    continue
        return None
    
    def _extract_year(self, text: str) -> Optional[int]:
        """Extract release/introduced year"""
        patterns = [
            # "Introduced January 2019", "Released January 2019", "Introduced: 2019"
            r'[•\-]?\s*(?:Introduced|Released):?\s+\w+\s+(\d{4})',
            # "Introduced December, 2006" (with comma)
            r'[•\-]?\s*(?:Introduced|Released):?\s+\w+,?\s+(\d{4})',
            # "Introduced 2019", "Released 2019", "Introduced: 2019"
            r'[•\-]?\s*(?:Introduced|Released):?\s+(\d{4})',
            # "Product Description 2012" (fallback pattern)
            r'Product\s+Description\s+(\d{4})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    year = int(match.group(1))
                    if 1990 <= year <= 2030:
                        return year
                except ValueError:
                    continue
        return None
    
    def _extract_retired_year(self, text: str) -> Optional[int]:
        """Extract retired year"""
        patterns = [
            # "Retired November 2021", "• Retired November 2021", "Retired: 2021"
            r'[•\-]?\s*Retired:?\s+\w+\s+(\d{4})',
            # "Retired December, 2008" (with comma)
            r'[•\-]?\s*Retired:?\s+\w+,?\s+(\d{4})',
            # "Retired 2021", "Retired: 2021"
            r'[•\-]?\s*Retired:?\s+(\d{4})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    year = int(match.group(1))
                    if 1990 <= year <= 2030:
                        return year
                except ValueError:
                    continue
        return None
    
    def _extract_purchased_year(self, text: str) -> Optional[int]:
        """Extract purchased year"""
        pattern = r'(?:Purchased?|Bought|Acquired)[\s:]+\w+\s+(\d{4})'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                return int(match.group(1))
            except ValueError:
                pass
        return None
    
    def _extract_collection(self, text: str) -> Optional[str]:
        """Extract collection name"""
        # Pattern: "Collection Name Collection" or "Part of the Collection Name Collection"
        patterns = [
            r'(?:Part of the|From the)\s+([^.]+?)\s+Collection',
            r'([^.]+?)\s+Collection',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                collection = match.group(1).strip()
                # Filter out generic phrases
                if collection and not re.match(r'^(a|an|the)$', collection, re.IGNORECASE):
                    return collection
        
        return None
    
    def _extract_series(self, text: str) -> Optional[str]:
        """Extract series name"""
        # Look for the specific line containing "Series"
        # Split by common separators to isolate just the series part
        for line in text.split('\n'):
            if 'Series' in line:
                # Pattern: match text immediately before " Series"
                match = re.search(r"([A-Z][A-Za-z\s']+)\s+Series", line)
                if match:
                    series = match.group(1).strip()
                    # Clean up: remove any SKU/Item number remnants
                    series = re.sub(r'(?:Box\s*\d*|Item\s*#?\s*\d+)\s*', '', series).strip()
                    # Filter out if it's too short or generic
                    if series and len(series) > 3 and not re.match(r'^(a|an|the)$', series, re.IGNORECASE):
                        return series
        
        return None
    
    def _extract_description(self, lines: List[str]) -> Optional[str]:
        """Extract description (first substantial sentence)"""
        for line in lines:
            # Skip lines that are just metadata
            if any(keyword in line.lower() for keyword in ['sku', 'item', 'introduced', 'retired', 'coordinates', '$', 'price']):
                continue
            
            # If line is long enough and looks like a description
            if len(line) > 30 and not line.isupper():
                return line
        
        return None
    
    def _extract_house_links(self, text: str, house_name: str) -> List[str]:
        """Extract links to houses"""
        patterns = [
            r'Coordinates\s+with\s+([^(]+?)(?:\s+\(|$)',
            r'Goes\s+with\s+([^(]+?)(?:\s+\(|$)',
        ]
        
        links = []
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                cleaned = match.strip()
                if cleaned and not self._names_match(cleaned, house_name):
                    links.append(cleaned)
        
        return links
    
    def _extract_notes(self, lines: List[str]) -> Optional[str]:
        """Extract notes - all text that's not SKU, price, dates, or description"""
        note_lines = []
        
        # Keywords that indicate metadata (not notes)
        metadata_keywords = ['sku', 'item #', 'top of form', 'bottom of form', '$']
        
        for line in lines:
            # Split by newlines in case there are embedded newlines
            sublines = line.split('\n')
            
            for subline in sublines:
                subline_stripped = subline.strip()
                
                # Skip empty lines
                if not subline_stripped:
                    continue
                
                subline_lower = subline_stripped.lower()
                
                # Skip standalone "Box" or "Box X" references
                if re.match(r'^\s*box\s*\d*\s*$', subline_lower):
                    continue
                
                # Skip metadata lines
                is_metadata = any(keyword in subline_lower for keyword in metadata_keywords)
                if is_metadata:
                    continue
                
                # Skip if it starts with introduced/retired/purchased/coordinates
                if re.match(r'^\s*(introduced|retired|purchased|coordinates)', subline_lower):
                    continue
                
                # Skip if it's just a price
                if re.match(r'^\$\s*\d+', subline_stripped):
                    continue
                
                # Skip if it's just a year
                if re.match(r'^\d{4}$', subline_stripped):
                    continue
                
                # Skip box lines
                if re.match(r'^Box\s+', subline_stripped, re.IGNORECASE):
                    continue
                
                # Add to notes if it passes all filters
                note_lines.append(subline_stripped)
        
        if note_lines:
            notes_text = '. '.join(note_lines)  # Join with periods for readability
            # Clean up common noise
            notes_text = re.sub(r'Imported from DOCX', '', notes_text, flags=re.IGNORECASE)
            notes_text = notes_text.strip()
            return notes_text if len(notes_text) > 5 else None
        
        return None
    
    def _is_box_line(self, line: str) -> bool:
        """Check if line is about box/packaging"""
        return bool(re.search(r'\bBox\s+\d+', line, re.IGNORECASE))
    
    def _is_no_accessory_marker(self, line: str) -> bool:
        """Check if line indicates there's no accessory"""
        no_acc_patterns = [
            r'No\s+ACC\s+Box',
            r'No\s+Accessory',
            r'No\s+ACC',
        ]
        
        for pattern in no_acc_patterns:
            if re.search(pattern, line, re.IGNORECASE):
                return True
        
        return False
    
    def _extract_accessory_list(self, line: str) -> List[tuple[str, Optional[str]]]:
        """
        Extract accessory names and descriptions from a line.
        Returns list of tuples: (name, description)
        Example: "Accessory - A Weekend Getaway – elf carrying trunk" 
                 -> [("A Weekend Getaway", "elf carrying trunk")]
        """
        # Check if line starts with "Accessories" or "Accessory"
        match = re.match(r'^\s*Accessor(?:y|ies)\s*[-:]?\s*(.+)', line, re.IGNORECASE)
        if not match:
            return []
        
        accessory_text = match.group(1).strip()
        
        # Clean up any trailing punctuation or form markers
        accessory_text = re.sub(r'(Bottom of Form|Top of Form).*$', '', accessory_text, flags=re.IGNORECASE).strip()
        
        # Split by "and" but ONLY when followed by a capital letter, Mr., Mrs., St., etc.
        # This keeps "and robe" with "Mrs. Clause with slippers" (lowercase 'r' in 'robe')
        # Pattern: split on "and" when followed by uppercase letter or title (case-sensitive!)
        parts = re.split(r'\s+and\s+(?=[A-Z]|Mrs?\.|St\.)', accessory_text)
        
        accessories = []
        for part in parts:
            part = part.strip()
            if not part:
                continue
            
            # Check if part contains " – " or " - " separating name from description
            # Look for em dash (–) or regular dash with title case before it
            name = part
            description = None
            
            # Try to split on em dash or regular dash
            for separator in ['–', ' - ', ' -']:
                if separator in part:
                    split_parts = part.split(separator, 1)
                    if len(split_parts) == 2:
                        potential_name = split_parts[0].strip()
                        potential_desc = split_parts[1].strip()
                        
                        # Check if the part before dash is title case (each word capitalized)
                        # This indicates it's the name, and what follows is description
                        words = potential_name.split()
                        if all(word[0].isupper() or word.lower() in ['a', 'an', 'the', 'of', 'for', 'to'] for word in words if word):
                            name = potential_name
                            description = potential_desc
                            break
            
            accessories.append((name, description))
        
        return accessories if accessories else []
    
    def _is_scenery_item(self, line: str) -> bool:
        """Check if line describes scenery items that should not be treated as accessories"""
        scenery_patterns = [
            r'^\d+\s+tree',  # "3 trees", "2 trees"
            r'^tree',        # "trees"
            r'^snow',        # "snow"
            r'^building$',   # "Building" (refers to house itself)
        ]
        
        line_lower = line.lower().strip()
        for pattern in scenery_patterns:
            if re.match(pattern, line_lower, re.IGNORECASE):
                return True
        
        return False
    
    def _extract_set_contents(self, paragraphs: List[str], start_idx: int) -> tuple[Optional[int], List[str]]:
        """
        Extract accessories from a 'Set of X' contents list.
        Returns: (index of 'Set of' line, list of accessory names)
        """
        set_idx = None
        
        # Find "Set of X" line
        for i in range(start_idx, len(paragraphs)):
            if re.match(r'^\s*Set\s+of\s+\w+', paragraphs[i], re.IGNORECASE):
                set_idx = i
                print(f"   Found set contents list at line {i}: {paragraphs[i]}")
                break
        
        if set_idx is None:
            return None, []
        
        # Extract items from following lines until we hit empty line or known ending marker
        accessories = []
        for i in range(set_idx + 1, len(paragraphs)):
            line = paragraphs[i].strip()
            
            # Stop at empty line or form markers
            if not line or re.search(r'(Bottom of Form|Top of Form)', line, re.IGNORECASE):
                break
            
            # Skip scenery items and house building reference
            if self._is_scenery_item(line):
                print(f"     Skipping scenery item: {line}")
                continue
            
            # This is an accessory
            print(f"     Found accessory from set: {line}")
            accessories.append(line)
        
        return set_idx, accessories
    
    def _names_match(self, name1: str, name2: str) -> bool:
        """Check if two names are similar enough to be the same"""
        # Simple similarity check
        name1 = name1.lower().strip()
        name2 = name2.lower().strip()
        
        # Exact match
        if name1 == name2:
            return True
        
        # One contains the other
        if name1 in name2 or name2 in name1:
            return True
        
        # Check if most words match
        words1 = set(name1.split())
        words2 = set(name2.split())
        if words1 and words2:
            overlap = len(words1 & words2)
            min_words = min(len(words1), len(words2))
            if overlap / min_words > 0.7:  # 70% word overlap
                return True
        
        return False
    
    def _extract_images(self, doc: Document, house: ExtractedItem, accessories: List[ExtractedItem], docx_path: str):
        """
        Extract images from the document
        - First image = house image (page 1)
        - Subsequent images = accessory images (page 2+)
        """
        if not house:
            return house, accessories
        
        print(f"\n{'─'*70}")
        print("Extracting images...")
        print(f"{'─'*70}")
        
        # Create images subdirectory
        images_dir = self.output_dir / "images"
        images_dir.mkdir(exist_ok=True)
        
        # Track images found
        images_found = []
        
        # Find all images in the document
        for para_idx, para in enumerate(doc.paragraphs):
            # Check if paragraph contains an image
            if para._element.xpath('.//pic:pic'):
                # Get the image relationship ID
                for rel_id in para._element.xpath('.//a:blip/@r:embed'):
                    rel = doc.part.rels[rel_id]
                    image_blob = rel.target_part.blob
                    
                    # Determine file extension from content type
                    content_type = rel.target_part.content_type
                    ext = 'jpg'
                    if 'png' in content_type:
                        ext = 'png'
                    elif 'jpeg' in content_type or 'jpg' in content_type:
                        ext = 'jpg'
                    elif 'gif' in content_type:
                        ext = 'gif'
                    
                    images_found.append({
                        'paragraph': para_idx,
                        'blob': image_blob,
                        'ext': ext,
                        'size': len(image_blob)
                    })
        
        print(f"Found {len(images_found)} image(s) in document")
        
        if not images_found:
            print("   ℹ️  No images found")
            return house, accessories
        
        # First image is the house
        if len(images_found) >= 1:
            house_image = images_found[0]
            # Create filename based on house name
            safe_name = re.sub(r'[^\w\s-]', '', house.name).strip().replace(' ', '_')
            filename = f"{safe_name}_house.{house_image['ext']}"
            filepath = images_dir / filename
            
            # Save image
            with open(filepath, 'wb') as f:
                f.write(house_image['blob'])
            
            # Update house record
            house.image_path = str(filepath)
            print(f"   📷 House image: {filename} ({house_image['size']:,} bytes)")
        
        # Remaining images are accessories
        if len(images_found) > 1:
            accessory_images = images_found[1:]
            
            # Match images to accessories by index
            for i, acc_image in enumerate(accessory_images):
                if i < len(accessories):
                    accessory = accessories[i]
                    # Create filename based on accessory name
                    safe_name = re.sub(r'[^\w\s-]', '', accessory.name).strip().replace(' ', '_')
                    filename = f"{safe_name}_accessory.{acc_image['ext']}"
                    filepath = images_dir / filename
                    
                    # Save image
                    with open(filepath, 'wb') as f:
                        f.write(acc_image['blob'])
                    
                    # Update accessory record
                    accessory.image_path = str(filepath)
                    print(f"   📷 Accessory image: {filename} ({acc_image['size']:,} bytes)")
                else:
                    print(f"   ⚠️  Extra image found but no matching accessory (image {i+2})")
        
        return house, accessories
    
    def save_results(self, results: Dict, output_filename: str):
        """Save parsing results to JSON"""
        output_path = self.output_dir / output_filename
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"\n✅ Results saved to: {output_path}")


def main():
    """Main function"""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docx_parser_simple.py <path_to_docx_file>")
        print("\nExample:")
        print('  python docx_parser_simple.py "\\\\DilgerNAS\\Public\\Media\\Day NP Files\\A Stitch in Yule Time 2019 She\'ll Be Belle of the Ball.docx"')
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    if not os.path.exists(docx_path):
        print(f"❌ Error: File not found: {docx_path}")
        sys.exit(1)
    
    parser = Dept56DocxParser()
    results = parser.parse_document(docx_path)
    
    # Display detailed results
    print("\n" + "="*70)
    print("EXTRACTED DATA FOR REVIEW")
    print("="*70)
    
    # Display houses
    for house in results['houses']:
        print("\n🏠 HOUSE")
        print("─"*70)
        print(f"  Name:           {house['name']}")
        print(f"  SKU:            {house['sku'] or '(not found)'}")
        print(f"  Description:    {house['description'] or '(not found)'}")
        print(f"  Price:          ${house['price']:.2f}" if house['price'] else "  Price:          (not found)")
        print(f"  Released Year:  {house['year'] or '(not found)'}")
        print(f"  Retired Year:   {house['retired_year'] or '(not found)'}")
        print(f"  Purchased Year: {house['purchased_year'] or '(not found)'}")
        print(f"  Collection:     {house['collection'] or '(not found)'}")
        print(f"  Series:         {house['series'] or '(not found)'}")
        print(f"  Linked Items:   {', '.join(house['linked_items']) if house['linked_items'] else '(none)'}")
        print(f"  Image:          {house['image_path'] or '(not found)'}")
        if house['notes']:
            print(f"  Notes:          {house['notes'][:200]}")
            if len(house['notes']) > 200:
                print(f"                  ... (truncated, see JSON for full notes)")
    
    # Display accessories
    for accessory in results['accessories']:
        print("\n🎨 ACCESSORY")
        print("─"*70)
        print(f"  Name:           {accessory['name']}")
        print(f"  SKU:            {accessory['sku'] or '(not found)'}")
        print(f"  Description:    {accessory['description'] or '(not found)'}")
        print(f"  Price:          ${accessory['price']:.2f}" if accessory['price'] else "  Price:          (not found)")
        print(f"  Released Year:  {accessory['year'] or '(not found)'}")
        print(f"  Retired Year:   {accessory['retired_year'] or '(not found)'}")
        print(f"  Purchased Year: {accessory['purchased_year'] or '(not found)'}")
        print(f"  Collection:     {accessory['collection'] or '(not found)'}")
        print(f"  Series:         {accessory['series'] or '(not found)'}")
        print(f"  Linked Items:   {', '.join(accessory['linked_items']) if accessory['linked_items'] else '(none)'}")
        print(f"  Image:          {accessory['image_path'] or '(not found)'}")
        if accessory['notes']:
            print(f"  Notes:          {accessory['notes'][:200]}")
            if len(accessory['notes']) > 200:
                print(f"                  ... (truncated, see JSON for full notes)")
    
    # Save results
    output_filename = f"parsed_{Path(docx_path).stem}.json"
    parser.save_results(results, output_filename)
    
    print("\n" + "="*70)
    print("✅ PARSING COMPLETE")
    print("="*70)
    print(f"\n📄 Full JSON saved to: {parser.output_dir / output_filename}")
    print(f"\n💡 Review the extracted data above!")
    print(f"   Let me know if fields are missing or incorrect.\n")


if __name__ == "__main__":
    main()
